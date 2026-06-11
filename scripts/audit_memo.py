"""Dump a generated memo's structure to eyeball formatting kinks."""
import sys
import pathlib

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parents[1] / "src"))
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document

path = sys.argv[1]
doc = Document(path)

print("### PARAGRAPHS (style | text) ###")
for p in doc.paragraphs:
    t = p.text.strip()
    if not t and "pBdr" not in p._p.xml:
        continue
    marker = "[RULED]" if "pBdr" in p._p.xml else ""
    style = p.style.name
    print(f"{style:<18} {marker}{t[:140]}")

print("\n### TABLES (first rows) ###")
for i, tb in enumerate(doc.tables):
    print(f"--- table {i}: {len(tb.rows)}x{len(tb.columns)}")
    for r in tb.rows[:3]:
        print("   | " + " | ".join(c.text.replace("\n", "\\n")[:40]
                                   for c in r.cells))

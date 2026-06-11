"""Render docs/GUIDE.md as a Word document (docs/CMA_Analyser_User_Guide.docx).

Handles the markdown subset the guide uses: #/##/### headings, tables,
fenced code blocks, bullet lists, bold/italic inline marks, hrules.
"""
import re
import sys
import pathlib

ROOT = pathlib.Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from docx import Document
from docx.shared import Pt, RGBColor

SRC = ROOT / "docs" / "GUIDE.md"
OUT = ROOT / "docs" / "CMA_Analyser_User_Guide.docx"


def add_runs(paragraph, text):
    """Inline **bold**, *italic* and `code` marks."""
    for part in re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def main():
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):                       # code block
            i += 1
            block = []
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(block))
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x57)
        elif line.startswith("|") and i + 1 < len(lines) \
                and set(lines[i + 1].replace("|", "").strip()) <= set("-: "):
            header = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            i -= 1
            table = doc.add_table(rows=0, cols=len(header))
            table.style = "Light Grid Accent 1"
            hdr = table.add_row().cells
            for c, text in zip(hdr, header):
                r = c.paragraphs[0].add_run(re.sub(r"\*+", "", text))
                r.bold = True
                r.font.size = Pt(9)
            for row in rows:
                cells = table.add_row().cells
                for c, text in zip(cells, row):
                    add_runs(c.paragraphs[0], text)
                    for p in c.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            add_runs(doc.add_paragraph(style="List Bullet"), line[2:])
        elif line.strip() == "---":
            pass
        elif line.strip():
            add_runs(doc.add_paragraph(), line)
        i += 1

    doc.save(OUT)
    print(f"Guide written: {OUT}")


if __name__ == "__main__":
    main()

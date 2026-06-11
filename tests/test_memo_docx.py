"""Tests for the credit memorandum Word export."""
import sqlite3
import pytest

from report.memo_docx import generate_credit_memo, DISCLAIMER
from ingest.demo_fixture import BORROWER


def _doc_text(path):
    from docx import Document
    doc = Document(path)
    paras = [p.text for p in doc.paragraphs]
    cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    return paras, cells, doc


def test_memo_generates_with_all_sections(ingested, tmp_path):
    data, bid, expected, db = ingested
    out = generate_credit_memo(BORROWER, out_path=tmp_path / "memo.docx",
                               db_path=db)
    assert out.exists()

    paras, cells, doc = _doc_text(out)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
                or p.style.name == "Title"]
    for expected_heading in (
            "CREDIT JUSTIFICATION MEMORANDUM",
            "1. Borrower Snapshot & Proposal",
            "2. Distress Score Ensemble",
            "3. Working Capital Assessment (Form V)",
            "4. Key Ratios vs Sanction Thresholds",
            "5. Red Flags, Exceptions & RBI EWS",
            "6. Projection Scrutiny (vs ETS bands)",
            "7. Questions for the Presenting Analyst",
            "8. Credit Committee Narrative (AI-drafted)"):
        assert any(expected_heading in h for h in headings), expected_heading

    # Mandatory disclaimer appears (top and bottom)
    assert sum(1 for p in paras if DISCLAIMER in p) >= 2
    # Form V content present with real numbers
    assert any("Working Capital Gap" in c for c in cells)
    assert any("MPBF" in c for c in cells)


def test_memo_includes_latest_narrative(ingested, tmp_path):
    data, bid, expected, db = ingested
    with sqlite3.connect(db) as conn:
        conn.execute("""
            INSERT INTO narrative (borrower_id, agent, model, prompt_hash, output_text)
            VALUES (?, 'underwriter', 'llama3.1:8b', 'abc123',
                    'RECOMMENDATION: APPROVE WITH COVENANTS based on DSCR.')
        """, (bid,))
    out = generate_credit_memo(BORROWER, out_path=tmp_path / "memo2.docx",
                               db_path=db)
    paras, _, _ = _doc_text(out)
    assert any("APPROVE WITH COVENANTS" in p for p in paras)


def test_memo_unknown_borrower_raises(temp_db, tmp_path):
    with pytest.raises(ValueError):
        generate_credit_memo("Ghost Co", out_path=tmp_path / "x.docx",
                             db_path=temp_db)

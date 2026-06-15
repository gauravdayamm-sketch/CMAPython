"""New units (projections only, no audited history) must assess on the
first projected year, clearly labelled — never blank."""
import csv
import io
import pytest

from ingest.llms_export import ingest_llms_set
from analytics.wc_assessment import assess_working_capital, form_v
from analytics.ews import run_full_ews
from analytics.scores import run_ensemble
from report.committee_queries import generate_queries
from report.memo_docx import generate_credit_memo

NAME = "GREENFIELD UNIT PVT LTD"

HEADER = [
    ["", f"Customer Id : C11111111  Customer Name :{NAME}"],
    [""],
    ["", "Accounting Year", "2026-27", "2027-28", "2028-29"],
    ["", "Financial Type", "Projected", "Projected", "Projected"],
]

BS = HEADER + [
    ["", "Current Liabilities: Short Term Borrowings from Banks"],
    ["1.", "From Applicant Bank",                       1.50, 1.80, 2.00],
    ["", "Current Liabilities: Others"],
    ["7.", "Sundry Creditors (Trade-Without LC)",       0.80, 0.95, 1.05],
    ["13.", "Deposits Instalments of Term Loans DPGs etc. (due within one year)",
                                                        0.40, 0.40, 0.40],
    ["", "Term Liabilities"],
    ["23.", "Secured Term Loans (excluding instalments payable within 1 year)",
                                                        3.20, 2.80, 2.40],
    ["", "Net Worth"],
    ["33.", "Equity Share Capital",                     2.00, 2.00, 2.00],
    ["37.", "Surplus/Deficit in Profit & Loss Account", 0.30, 0.85, 1.55],
    ["", "Current Assets"],
    ["45.", "Cash & Cash equivalents",                  0.25, 0.95, 1.75],
    ["", "Other Receivables"],
    ["50.", "Receivables (other than Deferred, Exports, BP & BD)",
                                                        1.40, 1.65, 1.85],
    ["", "Inventory"],
    ["56.", "Raw Material (Indigenous)",                1.10, 1.25, 1.40],
    ["", "Fixed Assets"],
    ["72.", "Gross Block",                              5.80, 5.80, 5.80],
    ["74.", "Depreciation (till date)",                 0.35, 0.85, 1.40],
]

OS = HEADER + [
    ["1.", "Domestic Sales",                            9.00, 11.00, 12.50],
    ["", "Raw Materials: (including Stores & other items used)"],
    ["9.", "Indigenous",                                6.30, 7.70, 8.75],
    ["14.", "Power & Fuel",                             0.45, 0.55, 0.62],
    ["17.", "Depreciation & Amortisation",              0.35, 0.50, 0.55],
    ["24.", "Selling & General Administrative Expenses", 0.70, 0.85, 0.95],
    ["26.", "Interest on Working Capital",              0.15, 0.18, 0.20],
    ["27.", "Interest on Term Loans",                   0.30, 0.26, 0.22],
    ["48.", "Provision for Taxes",                      0.18, 0.24, 0.30],
]


def _csv(rows):
    buf = io.StringIO()
    csv.writer(buf, lineterminator="\n").writerows(rows)
    return buf.getvalue()


@pytest.fixture
def new_unit(tmp_path, temp_db):
    bs = tmp_path / "G CMA_Balance_Sheet.csv"
    os_ = tmp_path / "G CMA_Operating_Statement.csv"
    bs.write_text(_csv(BS), encoding="utf-8")
    os_.write_text(_csv(OS), encoding="utf-8")
    data, bid = ingest_llms_set(bs, os_, db_path=temp_db)
    assert not data.errors()
    return temp_db, bid


def test_assessment_basis_is_first_projection(new_unit):
    db, bid = new_unit
    a = assess_working_capital(NAME, db_path=db)
    assert a.error == ""
    assert a.latest_audited is None
    year, label = a.assessment_basis
    assert year["fy_label"] == "2026-27"          # first projection
    assert "PROJECTED" in label and "new unit" in label
    rows = form_v(a)
    assert rows and rows[0][1] > 0                # TCA populated, not blank


def test_ews_runs_on_projection_basis(new_unit):
    db, bid = new_unit
    r = run_full_ews(NAME, db_path=db)
    assert r.error == ""
    assert "PROJECTED" in r.basis
    assert len(r.red_flags) == 32
    # Trend tests need a prior year and must degrade to N/A, not breach
    rf7 = next(f for f in r.red_flags if f.rule_id == "RF7")
    assert rf7.status == "N/A"


def test_ensemble_scores_projection_with_label(new_unit):
    db, bid = new_unit
    r = run_ensemble(NAME, db_path=db)
    assert r.error == ""
    assert "indicative" in r.basis
    assert r.altman is not None


def test_new_unit_query_generated(new_unit):
    db, bid = new_unit
    queries = generate_queries(NAME, db_path=db)
    nu = [q for q in queries if q["parameter"] == "New unit"]
    assert nu and nu[0]["priority"] == 1
    assert "independent evidence" in nu[0]["question"]


def test_memo_not_blank_for_new_unit(new_unit, tmp_path):
    db, bid = new_unit
    out = generate_credit_memo(NAME, out_path=tmp_path / "nu.docx", db_path=db)
    from docx import Document
    doc = Document(out)
    paras = [p.text for p in doc.paragraphs]
    cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    assert any("NEW UNIT" in p for p in paras)
    assert any("Working Capital Gap" in c for c in cells)
    # Form V values are populated, not dashes
    idx = next(i for i, c in enumerate(cells) if "Total Current Assets" in c)
    assert cells[idx + 1] not in ("", "—")
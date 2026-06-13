"""
Proposal-note consistency check (local-only).

Traces every figure quoted in the analyst's appraisal note (.docx) back to
the borrower's ingested CMA data. Figures that match are confirmed; the
rest are surfaced as verification items with their surrounding sentence —
typically facility limits and sanction terms (check against sanction
papers) or, occasionally, a number that exists nowhere.

Runs entirely on this machine; the note never leaves it.
"""
import pathlib
import re
import sqlite3

ROOT = pathlib.Path(__file__).resolve().parents[2]
DB   = ROOT / "db" / "cma.sqlite"

_NUM_RE = re.compile(r"\d[\d,]*(?:\.\d+)?")


def _docx_text(path):
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            parts.extend(c.text for c in r.cells)
    return "\n".join(parts)


def _known_values(borrower_name, db_path):
    """Every number the CMA data can vouch for, in common representations."""
    from analytics.wc_assessment import _load, compute_year_metrics
    borrower, years, thresholds, proposals = _load(borrower_name, db_path)
    if not borrower or not years:
        return None
    known = set()

    def add(v):
        if isinstance(v, (int, float)) and v is not None:
            v = float(v)
            for rep in (v, v * 100, v / 100):
                known.add(round(rep, 4))

    for y in years:
        for v in y["lines"].values():
            add(v)
        for v in compute_year_metrics(y["lines"]).values():
            if isinstance(v, (int, float)):
                add(v)
    for lim in proposals.values():
        add(lim.get("existing"))
        add(lim.get("proposed"))
    for v in thresholds.values():
        add(v)
    return known


def check_note(docx_path, borrower_name=None, db_path=None):
    """Returns {borrower, total, traced, untraced:[{figure,context}],
    trace_rate, error}."""
    db = db_path or DB
    if borrower_name is None:
        with sqlite3.connect(db) as conn:
            row = conn.execute("""
                SELECT b.name FROM borrower b
                JOIN cma_statement s ON s.borrower_id = b.borrower_id
                GROUP BY b.borrower_id ORDER BY MAX(s.imported_at) DESC LIMIT 1
            """).fetchone()
        borrower_name = row[0] if row else None
    known = _known_values(borrower_name, db) if borrower_name else None
    if known is None:
        return {"borrower": borrower_name, "total": 0, "traced": 0,
                "untraced": [], "trace_rate": None,
                "error": "No CMA data for this borrower"}

    text = _docx_text(docx_path)

    def traceable(n, raw):
        # Match at the precision the note quotes: a 2-dp figure must match
        # to ±0.005; a whole number tolerates ±0.5.
        decimals = min(len(raw.split(".")[1]) if "." in raw else 0, 2)
        tol = 0.5 * 10 ** (-decimals)
        return any(abs(n - v) <= tol for v in known)

    total = traced = 0
    untraced = []
    for m in _NUM_RE.finditer(text):
        raw = m.group(0)
        n = float(raw.replace(",", ""))
        before = text[m.start() - 1: m.start()]
        after = text[m.end(): m.end() + 2]
        if before.isalpha() or after[:1].isalpha():   # inside an identifier
            continue
        if before == "." or (after[:1] == "." and after[1:2].isdigit()):
            continue                                    # dotted chain / date
        if n == int(n) and (n < 10 or 1900 <= n <= 2100 or n >= 10000):
            continue                                    # index / year / code
        total += 1
        if traceable(n, raw):
            traced += 1
        else:
            start = max(0, m.start() - 60)
            untraced.append({"figure": raw,
                             "context": " ".join(text[start:m.end() + 60].split())})

    return {"borrower": borrower_name, "total": total, "traced": traced,
            "untraced": untraced,
            "trace_rate": round(traced / total, 3) if total else None,
            "error": ""}

"""
Credit Memorandum Word export — assembles everything the pipeline knows
about a borrower into a bank-format .docx:

  1. Borrower snapshot + proposal limits
  2. Distress score ensemble
  3. Form V working capital assessment (MPBF) + DSCR
  4. Key ratios vs thresholds
  5. Red flags / exception engine / RBI EWS
  6. Projection scrutiny
  7. AI committee narrative (latest from the narrative table), with the
     mandatory analyst-review disclaimer

Every quantitative section is computed fresh from the database at export
time — the document never goes stale relative to the data.
"""
import datetime
import pathlib
import sqlite3

ROOT = pathlib.Path(__file__).resolve().parents[2]
DB   = ROOT / "db" / "cma.sqlite"

DISCLAIMER = ("AI-DRAFTED DOCUMENT — REQUIRES ANALYST REVIEW AND SIGN-OFF "
              "BEFORE SUBMISSION.")

# Plain-language glossary for committee members without a credit background.
GLOSSARY = [
    ("CMA data", "Credit Monitoring Arrangement — the standard set of past and "
     "projected financial statements a borrower submits when seeking limits."),
    ("Working Capital Gap (WCG)", "Current assets minus current liabilities other "
     "than bank borrowings — the funding hole in the day-to-day business cycle "
     "that someone (the borrower or the bank) must fill."),
    ("NWC (Net Working Capital)", "The portion of current assets funded from the "
     "borrower's own long-term money. Negative NWC means short-term funds are "
     "propping up long-term uses — a classic stress signal."),
    ("Tandon Committee / Method I & II", "RBI committee (1975) that set how much "
     "of the working capital gap a borrower must fund itself. Method I: at least "
     "25% of the GAP from own funds. Method II (stricter, the standard): at least "
     "25% of ALL current assets from own funds."),
    ("MPBF", "Maximum Permissible Bank Finance — the most the bank should lend "
     "for working capital after the borrower's stipulated margin under the "
     "chosen Tandon method. A proposed limit above MPBF needs justification."),
    ("ABF", "Assessed Bank Finance — the gap actually left after the borrower's "
     "available margin (WCG minus actual NWC)."),
    ("Turnover (Nayak) method", "Simpler RBI yardstick for smaller limits: "
     "working capital need is taken as 25% of projected sales, of which the "
     "bank funds 20% of sales."),
    ("DSCR", "Debt Service Coverage Ratio — for every rupee of loan repayment "
     "due in a year (instalments + interest), how many rupees of cash the "
     "business generates. Below 1.00x the business cannot pay from its own "
     "cash. Gross DSCR includes interest in both numerator and denominator; "
     "Net DSCR is cash against instalments alone."),
    ("TNW", "Tangible Net Worth — the owners' real stake: capital plus "
     "accumulated profits, excluding intangibles."),
    ("TOL/TNW", "Total Outside Liabilities to Tangible Net Worth — how many "
     "rupees of other people's money sit on one rupee of the owners' money. "
     "Higher = more leveraged = riskier."),
    ("ICR", "Interest Coverage Ratio — operating profit divided by interest "
     "cost. Below ~1.5 the business barely earns its interest bill."),
    ("Debt/EBITDA", "Years of current cash earnings needed to repay term debt "
     "if everything else stayed still."),
    ("FACR", "Fixed Asset Coverage Ratio — value of fixed assets backing each "
     "rupee of term debt; the security cushion for term loans."),
    ("DSO / Inventory days / DPO", "How many days sales sit unpaid with "
     "customers / stock sits in the godown / the borrower takes to pay "
     "suppliers. Sudden stretching of any of these is an early warning."),
    ("Cash conversion cycle", "Inventory days + DSO − DPO: how long a rupee is "
     "stuck in the operating cycle before coming back as cash."),
    ("STBB / USL", "Short-Term Bank Borrowings; Unsecured Loans (typically from "
     "promoters or group entities — can be withdrawn overnight unless "
     "subordinated to bank debt)."),
    ("Dual column (Estimated vs Audited)", "The CMA shows the same year twice: "
     "what the borrower ESTIMATED a year ago and what was actually AUDITED. "
     "The gap measures how much the borrower's projections can be trusted."),
    ("ETS band", "A statistical forecast range fitted on the borrower's own "
     "audited history. A projection above the band needs evidence (orders, "
     "capacity), not optimism."),
    ("Ohlson / Altman Z″ / Zmijewski", "Three independent academic bankruptcy "
     "models (1980/1995/1984). Zones: SAFE / GREY / DISTRESS. Ohlson tends to "
     "read hot on Indian SMEs that run on creditor financing — agreement "
     "across models is the meaningful signal, not any single one."),
    ("EWS / RFA", "Early Warning Signals per RBI's Fraud Risk Master Direction "
     "(Jul 2024). RFA = Red Flagged Account — mandatory once any Critical "
     "indicator or two High indicators trigger; starts a fraud-review clock."),
]


def _note(doc, text):
    """Small italic explainer under a table, for non-credit readers."""
    from docx.shared import Pt
    p = doc.add_paragraph()
    run = p.add_run("Note: " + text)
    run.italic = True
    run.font.size = Pt(8)
    return p


def _ruled_line(doc):
    """An empty paragraph with a bottom border — a line to write on."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def _table(doc, rows, header=None, widths=None):
    from docx.shared import Pt
    n_cols = len(header) if header else len(rows[0])
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Light Grid Accent 1"
    if header:
        cells = table.add_row().cells
        for c, text in zip(cells, header):
            run = c.paragraphs[0].add_run(str(text))
            run.bold = True
            run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for c, value in zip(cells, row):
            run = c.paragraphs[0].add_run(
                "" if value is None else str(value))
            run.font.size = Pt(9)
    return table


def _fmt(v, pattern="{:,.2f}"):
    if v is None:
        return "—"
    if isinstance(v, (int, float)):
        return pattern.format(v)
    return str(v)


def _latest_narrative(borrower_name, db_path):
    with sqlite3.connect(db_path) as conn:
        conn.row_factory = sqlite3.Row
        row = conn.execute("""
            SELECT n.output_text, n.model, n.run_ts FROM narrative n
            JOIN borrower b ON b.borrower_id = n.borrower_id
            WHERE b.name = ? AND n.agent = 'underwriter'
            ORDER BY n.run_ts DESC, n.narr_id DESC LIMIT 1
        """, (borrower_name,)).fetchone()
    return row


def generate_credit_memo(borrower_name=None, out_path=None, db_path=None):
    """Build the .docx. Returns the output path."""
    from docx import Document
    from docx.shared import Pt

    from analytics.wc_assessment import assess_working_capital, form_v
    from analytics.ews import run_full_ews
    from analytics.scores import run_ensemble
    from analytics.forecast import scrutinize_projections

    db = db_path or DB
    a = assess_working_capital(borrower_name, db_path=db)
    if a.error:
        raise ValueError(a.error)
    name = a.borrower_name

    ews = run_full_ews(name, db_path=db)
    ens = run_ensemble(name, db_path=db)
    scr = scrutinize_projections(name, db_path=db)

    out_path = pathlib.Path(out_path or ROOT / "data" /
                            f"Credit_Memo_{name.replace(' ', '_')}.docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("CREDIT JUSTIFICATION MEMORANDUM", level=0)
    p = doc.add_paragraph()
    p.add_run(f"{name}\n").bold = True
    p.add_run(f"Generated: {datetime.datetime.now():%d %B %Y %H:%M}   |   "
              f"CMA Auto-Analytics pipeline")
    doc.add_paragraph(DISCLAIMER).runs[0].bold = True

    # 1 ── Snapshot
    _heading(doc, "1. Borrower Snapshot & Proposal")
    latest = a.latest_audited
    snap = [
        ("Latest audited FY", latest["fy_label"] if latest else "—"),
        ("Net sales (₹ Cr)", _fmt(latest["metrics"]["net_sales"]) if latest else "—"),
        ("TNW (₹ Cr)", _fmt(latest["metrics"]["tnw"]) if latest else "—"),
        ("PAT (₹ Cr)", _fmt(latest["metrics"]["pat"]) if latest else "—"),
    ]
    for facility, lim in a.proposals.items():
        snap.append((f"Limit — {facility}",
                     f"existing {_fmt(lim['existing'])} / "
                     f"proposed {_fmt(lim['proposed'])}"))
    _table(doc, snap, header=("Item", "Value"))

    # 2 ── Score ensemble
    _heading(doc, "2. Distress Score Ensemble")
    if not ens.error:
        rows = []
        for s in (ens.ohlson, ens.altman, ens.zmijewski):
            if s:
                rows.append((s.model, _fmt(s.score, "{:.4f}"),
                             f"{s.prob:.1%}" if s.prob is not None else "—",
                             s.zone, s.note or ""))
        _table(doc, rows, header=("Model", "Score", "Prob.", "Zone", "Note"))
        doc.add_paragraph(f"Ensemble verdict: {ens.verdict}")
        _note(doc, "Three independent bankruptcy models from different eras "
                   "and methods. No single model decides; agreement across "
                   "them is the signal. Ohlson typically reads hot on Indian "
                   "SMEs that run on creditor financing. See Annexure A.")
    else:
        doc.add_paragraph(f"Not available: {ens.error}")

    # 3 ── Working capital + DSCR
    _heading(doc, "3. Working Capital Assessment (Form V)")
    _table(doc, [(label, _fmt(value)) for label, value in form_v(a)],
           header=("Form V item", "₹ Cr"))
    _note(doc, "Reading this table: the Working Capital Gap (row 3) is the "
               "funding hole in the operating cycle. The RBI's Tandon "
               "committee requires the borrower to fill part of it from own "
               "long-term funds — Method I: 25% of the gap; Method II "
               "(stricter, standard): 25% of all current assets. MPBF "
               "(rows 7–8) is the most the bank should lend after that "
               "margin; a proposed limit above MPBF needs explicit "
               "justification. ABF (row 9) is the gap left after the "
               "borrower's actual margin. See Annexure A.")
    doc.add_paragraph()
    dscr_rows = [
        (y["fy_label"], y["statement_type"],
         _fmt(y["metrics"]["gross_dscr"], "{:.2f}"),
         _fmt(y["metrics"]["net_dscr"], "{:.2f}"))
        for y in a.years
        if y["metrics"]["gross_dscr"] is not None and not y["is_dual"]
    ]
    if dscr_rows:
        _table(doc, dscr_rows, header=("FY", "Type", "Gross DSCR", "Net DSCR"))
        doc.add_paragraph(f"Average gross DSCR: {_fmt(a.dscr_avg_gross, '{:.2f}')}x")
        _note(doc, "DSCR: for every rupee of loan repayment due in the year "
                   "(instalments + interest), how many rupees of cash the "
                   "business generates. Below 1.00x it cannot service debt "
                   "from its own cash; the sanction benchmark here is "
                   f"{a.thresholds.get('min_dscr_any', 1.25)}x in every year "
                   f"and {a.thresholds.get('min_dscr_avg', 1.5)}x on average.")

    # 4 ── Ratios vs thresholds
    _heading(doc, "4. Key Ratios vs Sanction Thresholds")
    breaches = [v for v in a.verdicts if v["status"] == "BREACH"]
    latest_verdicts = [v for v in a.verdicts
                       if latest and v["fy_label"] == latest["fy_label"]]
    _table(doc,
           [(v["metric"], _fmt(v["value"], "{:.2f}"),
             f"{v['direction']} {v['threshold']}", v["status"])
            for v in latest_verdicts],
           header=("Ratio (latest audited)", "Value", "Threshold", "Status"))
    doc.add_paragraph(
        f"Threshold breaches across all years: {len(breaches)}")
    _note(doc, "TOL/TNW = rupees of outside money per rupee of the owners' "
               "money (higher = more leveraged). ICR = operating profit vs "
               "the interest bill. FACR = fixed assets backing each rupee of "
               "term debt. Day-count ratios (DSO, inventory days, DPO) show "
               "where cash is stuck in the cycle. Annexure A explains each.")

    # 5 ── Red flags / EWS
    _heading(doc, "5. Red Flags, Exceptions & RBI EWS")
    if not ews.error:
        doc.add_paragraph(f"Red flag verdict: {ews.red_flag_verdict}")
        doc.add_paragraph(
            "RFA (Red Flagged Account): "
            + (f"REQUIRED — {ews.rfa_reason}" if ews.rfa_required
               else "not indicated"))
        hits = ews.triggered_red_flags + ews.triggered_exceptions
        if hits:
            _table(doc,
                   [(f.rule_id, f.test, _fmt(f.value, "{:.4f}"), f.detail[:120])
                    for f in hits],
                   header=("Rule", "Test", "Value", "Action required"))
        ews_hits = [i for i in ews.ews_indicators if i["triggered"]]
        if ews_hits:
            _table(doc,
                   [(f"#{i['id']}", i["indicator"], i["severity"], i["source"])
                    for i in ews_hits],
                   header=("EWS", "Indicator", "Severity", "Source"))
    else:
        doc.add_paragraph(f"Not available: {ews.error}")

    # 5b ── Market & registry intelligence
    try:
        from data.registry_checks import get_checks
        from data.news_intel import get_news
        with sqlite3.connect(db) as conn:
            bid_row = conn.execute(
                "SELECT borrower_id FROM borrower WHERE name = ?",
                (name,)).fetchone()
        news = get_news(bid_row[0], db_path=db) if bid_row else []
        adverse = [n for n in news if n["classification"] == "ADVERSE"]
        doc.add_paragraph()
        doc.add_paragraph(
            f"Market intelligence: {len(adverse)} adverse / {len(news)} "
            f"headlines on record.")
        for n in adverse[:5]:
            doc.add_paragraph(
                f"⚠ {n['title'][:110]} ({n['published'][:10]})",
                style="List Bullet")
        checks = get_checks(name, db_path=db)
        if checks:
            _table(doc,
                   [(c["label"], c["status"].upper(),
                     c["checked_on"] or "—", (c["remarks"] or "")[:60])
                    for c in checks],
                   header=("Registry due diligence", "Status",
                           "Checked on", "Remarks"))
            _note(doc, "Registries marked UNCHECKED or STALE (>90 days) "
                       "appear as committee queries in section 7.")
    except Exception:
        pass

    # 6 ── Projection scrutiny
    _heading(doc, "6. Projection Scrutiny (vs ETS bands)")
    if not scr["error"]:
        for metric, entry in scr["metrics"].items():
            if entry["error"]:
                doc.add_paragraph(f"{metric}: {entry['error']}")
                continue
            if entry["hist_cagr"] is not None and entry["proj_cagr"] is not None:
                cagr = (f" — historical CAGR {entry['hist_cagr']:.1%}, "
                        f"projected {entry['proj_cagr']:.1%}")
            else:
                cagr = " — growth trend not computable (loss years in history)"
            doc.add_paragraph(f"{metric}{cagr}", style="Intense Quote")
            _table(doc,
                   [(pp["fy_label"], _fmt(pp["value"]),
                     f"{_fmt(pp['lo_80'])} – {_fmt(pp['hi_80'])}", pp["verdict"])
                    for pp in entry["projections"]],
                   header=("FY", "Projected", "ETS 80% band", "Verdict"))
    else:
        doc.add_paragraph(f"Not available: {scr['error']}")

    # 7 ── Questions for the presenting analyst
    from report.committee_queries import generate_queries
    _heading(doc, "7. Questions for the Presenting Analyst")
    queries = generate_queries(name, db_path=db)
    if queries:
        doc.add_paragraph(
            "Raised automatically from breached parameters, red flags and "
            "optimistic projections — each question cites the figure that "
            "prompted it. Priority 1 items go to the analyst first.")
        for i, q in enumerate(queries, start=1):
            p = doc.add_paragraph()
            p.add_run(f"Q{i} [P{q['priority']}] {q['parameter']}: ").bold = True
            p.add_run(q["observation"])
            doc.add_paragraph(q["question"], style="Intense Quote")
    else:
        doc.add_paragraph("No adverse findings — no committee queries generated.")

    # 8 ── Committee narrative
    _heading(doc, "8. Credit Committee Narrative (AI-drafted)")
    narr = _latest_narrative(name, db)
    if narr:
        doc.add_paragraph(
            f"Drafted by {narr['model']} on {narr['run_ts']} — full agent "
            f"attribution in the narrative audit table.")
        for para in narr["output_text"].split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    else:
        doc.add_paragraph(
            "No committee narrative on record — run the 4-agent committee "
            "(POST /narrative) and re-export.")

    # 9 ── Committee member's own remarks
    _heading(doc, "9. Committee Member's Remarks")
    doc.add_paragraph(
        "Observations, additional queries raised in the meeting, and "
        "conditions to be stipulated:")
    for _ in range(8):
        _ruled_line(doc)
    doc.add_paragraph()
    doc.add_paragraph(
        "Recommendation:   ☐ Approve      ☐ Approve with conditions      "
        "☐ Defer (information sought)      ☐ Decline")
    sig = _table(doc, [("", "", "")],
                 header=("Name & Designation", "Signature", "Date"))
    for cell in sig.rows[1].cells:
        cell.add_paragraph()
        cell.add_paragraph()

    # Annexure ── plain-language glossary
    _heading(doc, "Annexure A — Glossary for Committee Members")
    doc.add_paragraph(
        "Plain-language explanations of the technical terms used in this "
        "memorandum, for members without a credit background.")
    _table(doc, GLOSSARY, header=("Term", "What it means"))

    doc.add_paragraph()
    doc.add_paragraph(DISCLAIMER).runs[0].bold = True

    try:
        doc.save(out_path)
    except PermissionError:
        # Target is open in Word — save a timestamped copy instead
        stamped = out_path.with_name(
            f"{out_path.stem}_{datetime.datetime.now():%H%M%S}{out_path.suffix}")
        doc.save(stamped)
        return stamped
    return out_path
